# Create Word (.docx) for Well Pump Solar ELECTRICAL INSTALL ONLY — run from Well Pump folder
from docx import Document
from docx.shared import Pt

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    return p

def main():
    import os
    folder = os.path.dirname(os.path.abspath(__file__))

    doc = Document()
    doc.add_paragraph('Request for Proposal', style='Title')
    doc.add_paragraph('Well Pump Solar Power System — Electrical Installation Only (Solar Electrical Vault)', style='Heading 1')
    doc.add_paragraph()
    add_para(doc, 'Project: Rodriguez Residence')
    add_para(doc, 'Location: Oak Valley Estates Lot 5, Kanarraville, UT — Well vault and fence-mounted solar array')
    add_para(doc, 'Delivery: Lump sum for electrical labor per attached scope')
    add_para(doc, 'Reference: All equipment Owner-furnished. See Well Pump\\HTML\\4_Well_Solar_Visual.html and related docs in the Well Pump folder.')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('Definition — Solar Electrical Vault Scope', style='Heading 2')
    doc.add_paragraph('Vault: Underground enclosure, 36" x 42", hatch at grade. Contains Victron MPPT, 24V battery bank (100Ah x2 series), Victron MultiPlus-II 24/3000 2x120V. DC in from solar; AC out to VFD.', style='List Bullet')
    doc.add_paragraph('Already installed (do not touch): 705W solar panel (fence-mounted); Franklin pump, pump wire, VFD Commander Pro at wellhead.', style='List Bullet')
    doc.add_paragraph('Contractor electrical scope: (1) PV combiner at array (2) DC cable array to vault (3) Vault interior wiring (4) AC cable vault to VFD AC input only (5) Commissioning and labeling.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('1. Scope of Work — Electrical Installation Only', style='Heading 2')
    add_para(doc, 'Proposals are for ELECTRICAL INSTALLATION ONLY. No structural, fence, or concrete. All equipment Owner-furnished. Contractor wires per reference.', bold=False)
    doc.add_paragraph('Flow: 705W panel (already up) → PV Combiner at array → DC cable to vault → Victron MPPT → 24V battery bank → Victron MultiPlus-II 24/3000 → AC to VFD. VFD and pump already installed — AC input terminals only.', style='List Bullet')
    doc.add_paragraph('PV Combiner at array: Mount combiner (fuse, 40A DC breaker, SPD) near panel. MC4 in, DC out to vault. Bond ground to array and GEC.', style='List Bullet')
    doc.add_paragraph('DC cable: Run USE-2/PV in conduit from combiner to vault. Terminate at charge controller.', style='List Bullet')
    doc.add_paragraph('Vault: Mount MPPT; wire PV in, battery out. Place batteries (24V, 100Ah x2 series); fuse; connect to controller and inverter DC. Mount MultiPlus-II; DC from battery, AC L1/L2 to run to VFD. Sequence: batteries to controller and inverter first, then PV to controller. Never PV before batteries.', style='List Bullet')
    doc.add_paragraph('AC cable: Vault to VFD Commander Pro AC input only (L1, L2, ground). 240V. Do not open VFD.', style='List Bullet')
    doc.add_paragraph('Commissioning: Batteries on → controller PV reading → inverter on → confirm VFD power. Do not test pump. Document string voltage/current if requested. Label all.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('2. Owner-Furnished', style='Heading 2')
    doc.add_paragraph('705W panel (installed). Franklin pump and VFD (installed). Combiner, Victron MPPT, battery bank, MultiPlus-II 24/3000, cable, conduit, fuses, misc. per reference. Vault as existing.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('3. Contractor Provides', style='Heading 2')
    doc.add_paragraph('Labor: combiner mount, DC run, vault wiring per sequence, AC run to VFD, terminate at AC input only, commissioning, labeling. Tools and test equipment. Confirm end-to-end operation.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('4. Pricing', style='Heading 2')
    doc.add_paragraph('Lump sum for electrical install only. Payment terms. Exclusions or assumptions.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('5. Proposal Format and Due Date', style='Heading 2')
    doc.add_paragraph('(1) Cover letter — company, contact. (2) Confirmation scope is electrical only per reference. (3) Lump sum, payment terms, exclusions. (4) Availability/lead time. (5) Questions.', style='List Number')
    add_para(doc, 'Submit to: [Insert contact name/email]')
    add_para(doc, 'Due date: [Insert date]')
    doc.add_paragraph()
    add_para(doc, 'This RFP is for electrical installation labor only. Equipment and design are in the Well Pump folder (e.g. 4_Well_Solar_Visual.html). Panel and VFD/pump already in place.', bold=False)
    well_pump_folder = os.path.dirname(folder)  # Well Pump root
    out = os.path.join(well_pump_folder, 'RFP - Well Pump Solar Electrical Install.docx')
    doc.save(out)
    print('Created:', out)

if __name__ == '__main__':
    main()
