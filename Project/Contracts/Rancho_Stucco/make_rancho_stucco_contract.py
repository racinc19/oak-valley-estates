# Creates Rancho Stucco LLC Subcontract (Rodriguez Residence)
# Run: python make_rancho_stucco_contract.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()
doc.add_paragraph("SUBCONTRACT AGREEMENT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("STUCCO — EXTERIOR").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence  |  Oak Valley Estates Lot 5  |  Kanarraville, UT 84742")
doc.add_paragraph()

add_heading(doc, "1. Scope of Work")
add_para(doc, "Subcontractor shall furnish all labor, materials, equipment, and delivery necessary to perform complete exterior stucco services for the Rodriguez Residence (Oak Valley Estates #5) per the Contract Documents and per Estimate No. 1947 dated 09/18/2025. Scope includes: water restrictive sheets; 1\" foam boards; 20 gauge metal wire; cement base coat; synthetic color coat; patio lids; soffits/fascias; scratch coat; labor; materials; and delivery.")
add_para(doc, "Total Contract Amount: $45,000.00 (Forty-Five Thousand Dollars).")

add_heading(doc, "2. Exclusions")
add_para(doc, "Any work not specifically described in this agreement shall be considered extra and requires written authorization from the General Contractor prior to commencement.")

add_heading(doc, "3. Contract Amount and Payment")
add_para(doc, "Total Contract Amount: $45,000.00 (Forty-Five Thousand Dollars).")
add_para(doc, "Subcontractor shall submit progress invoices as agreed (e.g., monthly or by phase). Each invoice shall describe work completed and applicable percentage of contract value. General Contractor shall review, approve, and pay within thirty (30) days of receipt. A retention percentage may be agreed in writing; if not specified, ten percent (10%) may be withheld until final completion and punchlist sign-off.")
add_para(doc, "Invoices shall be submitted in writing (email acceptable) to: racinc19@gmail.com. Include: (1) invoice number and date, (2) billing period, (3) description of work completed, (4) percentage complete, (5) amount billed this period, (6) total billed to date.")

add_heading(doc, "4. Schedule")
add_para(doc, "Subcontractor shall commence work upon written Notice to Proceed from the General Contractor. Subcontractor shall maintain continuous progress as coordinated with the project schedule. Any delays caused by the Subcontractor may result in liquidated or actual damages at the General Contractor's election.")

add_heading(doc, "5. Changes and Change Orders")
add_para(doc, "No extra work or changes to the Scope of Work shall be performed without prior written authorization from the General Contractor. Work performed without written approval will not be compensated. Subcontractor shall submit a written cost proposal for any change within five (5) business days of request.")

add_heading(doc, "6. Site Conditions and Responsibilities")
add_para(doc, "Subcontractor shall verify substrate and dimensions in the field. Discrepancies must be reported in writing immediately. Daily cleanup required. Subcontractor shall comply with all applicable OSHA regulations and maintain a safe worksite.")

add_heading(doc, "7. Insurance Requirements")
add_para(doc, "Prior to commencing work, Subcontractor shall obtain and maintain: Commercial General Liability: $1,000,000 per occurrence / $2,000,000 aggregate; Workers' Compensation as required by Utah state law. RAC Inc. and ARJR Kanarraville Holdings, LLC shall be named as additional insureds on the CGL policy. Certificates of Insurance must be delivered to General Contractor before mobilization.")

add_heading(doc, "8. Indemnification")
add_para(doc, "Subcontractor shall indemnify, defend, and hold harmless RAC Inc., ARJR Kanarraville Holdings, LLC, and their officers, employees, and agents from and against any and all claims, damages, losses, and expenses, including attorney's fees, arising out of or resulting from the Subcontractor's performance under this Agreement, to the extent caused by the negligent acts or omissions of the Subcontractor or its employees.")

add_heading(doc, "9. Termination")
add_para(doc, "General Contractor may terminate for cause upon written notice if Subcontractor: (a) abandons the work; (b) fails to maintain required insurance; (c) performs defective work and fails to remedy within five (5) days of notice; or (d) becomes insolvent. General Contractor may terminate for convenience upon seven (7) days' written notice, in which case Subcontractor shall be compensated for work satisfactorily completed to the date of termination.")

add_heading(doc, "10. Subcontractor Login Information (Project Master List)")
add_para(doc, "Project vendor and login access are tracked by Category, Vendor, and Pin per the project master list. General Contractor shall use this information only for project administration, document access, and communication and shall keep it confidential.")
doc.add_paragraph("Category: Exterior Stucco")
doc.add_paragraph("Vendor: Rancho Stucco LLC")
doc.add_paragraph("Pin: 4440")
doc.add_paragraph()

doc.add_paragraph("Contract Date | _____________, ______")
doc.add_paragraph("Owner | ARJR Kanarraville Holdings, LLC")
doc.add_paragraph("General Contractor | RAC Inc.  |  Craig Engel  |  1267 N. 1390 W., St. George, UT 84770  |  435-229-7360")
doc.add_paragraph("Subcontractor | Rancho Stucco LLC  |  PO Box 790195, Virgin, UT 84779  |  ranchostuccollc@gmail.com")
doc.add_paragraph("Contract Amount | $45,000.00 (Forty-Five Thousand Dollars)  |  Estimate No. 1947 (09/18/2025)")
doc.add_paragraph("Start Date | Upon NTP")
doc.add_paragraph()

doc.add_paragraph("GENERAL CONTRACTOR:")
doc.add_paragraph("RAC Inc.")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Craig Engel")
doc.add_paragraph("Title: General Contractor")
doc.add_paragraph("Date: _____________________________")

doc.add_paragraph()
doc.add_paragraph("SUBCONTRACTOR:")
doc.add_paragraph("Rancho Stucco LLC")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: _____________________________")
doc.add_paragraph("Title: _____________________________")
doc.add_paragraph("Date: _____________________________")

out = "Rancho_Stucco_Subcontract_Rodriguez.docx"
doc.save(out)
print("Saved:", out)
