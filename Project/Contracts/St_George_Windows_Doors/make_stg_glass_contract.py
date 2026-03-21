# Creates St George Windows, Doors & More LLC Subcontract (Rodriguez Residence)
# Run: python make_stg_glass_contract.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()
doc.add_paragraph("SUBCONTRACT AGREEMENT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("ALUMINUM WINDOWS & DOORS — GLASS").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence  |  Oak Valley Estates Lot 5  |  Kanarraville, UT 84742")
doc.add_paragraph()

add_heading(doc, "1. Scope of Work")
add_para(doc, "Subcontractor shall furnish all materials and labor necessary to supply and install Strata thermally broken aluminum windows and doors for the Rodriguez Residence (Oak Valley Estates Lot 5) per Estimate No. 3692 dated 10/30/2025 and the Contract Documents. Scope includes: Strata Windows & Doors (thermally broken aluminum) per line items; operable, fixed, fixed XL, open-to-below (O2B) fixed and O2B XL, fixed corner, and curtain wall window installation at quoted rates per square meter; Strata thermally broken aluminum sliding doors (multiple configurations and sizes per estimate); glass suction trolley; forklift and boom lift rental; window silicone sausage; sticky window flashing; and all other items set forth in Estimate No. 3692.")
add_para(doc, "Total Contract Amount: $290,385.95 (Two Hundred Ninety Thousand Three Hundred Eighty-Five Dollars and Ninety-Five Cents).")

add_heading(doc, "2. Exclusions")
add_para(doc, "Any work or materials not specifically described in this agreement or in Estimate No. 3692 shall be considered extra and requires written authorization from the General Contractor prior to commencement.")

add_heading(doc, "3. Contract Amount and Payment")
add_para(doc, "Total Contract Amount: $290,385.95 (Two Hundred Ninety Thousand Three Hundred Eighty-Five Dollars and Ninety-Five Cents).")
add_para(doc, "Payment terms per Estimate No. 3692: 65% due to order; 25% due at or prior to delivery; 10% due upon completion. Subcontractor shall submit invoices in accordance with these milestones. General Contractor shall pay within the agreed schedule. Past-due balances may accrue interest at 2% per month (24% APR) and collection charges as stated in Subcontractor's terms.")
add_para(doc, "Invoices shall be submitted in writing (email acceptable) to: racinc19@gmail.com. Include invoice number, date, description of work/materials, and amount.")

add_heading(doc, "4. Schedule")
add_para(doc, "Subcontractor shall commence work and order materials upon written Notice to Proceed and receipt of initial payment as agreed. Subcontractor shall perform in accordance with the project schedule. Any delays caused by the Subcontractor may result in liquidated or actual damages at the General Contractor's election. Estimate expires 12/01/2025; contract execution confirms pricing.")

add_heading(doc, "5. Changes and Change Orders")
add_para(doc, "No extra work or changes to the Scope of Work shall be performed without prior written authorization from the General Contractor. Work or materials supplied without written approval will not be compensated. Subcontractor shall submit a written cost proposal for any change within five (5) business days of request.")

add_heading(doc, "6. Site Conditions and Responsibilities")
add_para(doc, "Subcontractor shall verify openings and conditions in the field. Discrepancies must be reported in writing immediately. Subcontractor shall comply with all applicable OSHA regulations and maintain a safe worksite. By accepting this agreement, Customer allows Subcontractor to order items listed and perform applicable labor; some items may not be returned or restocked.")

add_heading(doc, "7. Insurance Requirements")
add_para(doc, "Prior to commencing work, Subcontractor shall obtain and maintain: Commercial General Liability: $1,000,000 per occurrence / $2,000,000 aggregate; Workers' Compensation as required by Utah state law. RAC Inc. and ARJR Kanarraville Holdings, LLC shall be named as additional insureds on the CGL policy. Certificates of Insurance must be delivered to General Contractor before mobilization.")

add_heading(doc, "8. Indemnification")
add_para(doc, "Subcontractor shall indemnify, defend, and hold harmless RAC Inc., ARJR Kanarraville Holdings, LLC, and their officers, employees, and agents from and against any and all claims, damages, losses, and expenses, including attorney's fees, arising out of or resulting from the Subcontractor's performance under this Agreement, to the extent caused by the negligent acts or omissions of the Subcontractor or its employees.")

add_heading(doc, "9. Termination")
add_para(doc, "General Contractor may terminate for cause upon written notice if Subcontractor: (a) abandons the work; (b) fails to maintain required insurance; (c) performs defective work and fails to remedy within five (5) days of notice; or (d) becomes insolvent. General Contractor may terminate for convenience upon seven (7) days' written notice, in which case Subcontractor shall be compensated for work and materials satisfactorily provided to the date of termination.")

add_heading(doc, "10. Subcontractor Login Information (Project Master List)")
add_para(doc, "Project vendor and login access are tracked by Category, Vendor, and Pin per the project master list. General Contractor shall use this information only for project administration, document access, and communication and shall keep it confidential.")
doc.add_paragraph("Category: Exterior Glass")
doc.add_paragraph("Vendor: St George Windows, Doors & More LLC")
doc.add_paragraph("Pin: _________________________")
doc.add_paragraph()

doc.add_paragraph("Contract Date | _____________, ______")
doc.add_paragraph("Owner | ARJR Kanarraville Holdings, LLC")
doc.add_paragraph("General Contractor | RAC Inc.  |  Craig Engel  |  1267 N. 1390 W., St. George, UT 84770  |  435-229-7360")
doc.add_paragraph("Subcontractor | St George Windows, Doors & More LLC  |  5505 W 290 N, Hurricane, UT 84737  |  435-673-7377  |  billing@stgwindows.com")
doc.add_paragraph("Contract Amount | $290,385.95 (Two Hundred Ninety Thousand Three Hundred Eighty-Five Dollars and Ninety-Five Cents)  |  Estimate No. 3692 (10/30/2025)")
doc.add_paragraph("Start Date | Upon NTP and initial payment")
doc.add_paragraph()

doc.add_paragraph("GENERAL CONTRACTOR:")
doc.add_paragraph("RAC Inc.")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Craig Engel")
doc.add_paragraph("Title: General Contractor")
doc.add_paragraph("Date: _____________________________")

doc.add_paragraph()
doc.add_paragraph("SUBCONTRACTOR:")
doc.add_paragraph("St George Windows, Doors & More LLC")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: _____________________________")
doc.add_paragraph("Title: _____________________________")
doc.add_paragraph("Date: _____________________________")

out = "St_George_Windows_Doors_Subcontract_Rodriguez.docx"
doc.save(out)
print("Saved:", out)
