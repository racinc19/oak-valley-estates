# Creates Oliver Finishing Clayworks Subcontract (Rodriguez Residence)
# Run: python make_oliver_clayworks_contract.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()
doc.add_paragraph("SUBCONTRACT AGREEMENT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("CLAYWORKS / PLASTER INSTALLATION — LABOR").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence  |  Oak Valley Estates Lot 5  |  Kanarraville, UT 84742")
doc.add_paragraph()

add_heading(doc, "1. Scope of Work")
add_para(doc, "Subcontractor shall furnish all labor, supervision, tools, equipment, and materials necessary to perform Clayworks (rammed earth / lime-based plaster) installation for the Rodriguez Residence per the Contract Documents and design intent. Scope includes rammed earth finish interior with lime-based plaster, rammed earth finish exterior labor, and Clayworks material as set forth in the Schedule of Values below. The stated totals include logistics, additional protection, one-use tools, per diem, overtime or minor delays due to weather or illness, taxes, and insurance coverage.")
add_para(doc, "Schedule of Values (Estimate No. 60, dated 02/20/2026):")
add_para(doc, "• Rammed earth finish interior with lime-based plaster: $40.00/sf × 745 sf = $29,800.00\n• Rammed earth finish exterior labor: $45.00/sf × 830 sf = $37,350.00\n• Clayworks material: $18,000.00\n• Hotel/Airbnb (est., up to 1 month for 2 people; may change): $3,500.00\n• Car rental (prices may vary): $1,500.00\n• Flights (prices may vary): $1,500.00")
add_para(doc, "Total Contract Amount: $91,650.00 (Ninety-One Thousand Six Hundred Fifty Dollars). Travel/lodging line items are estimates and may be adjusted by mutual written agreement if actual costs differ.")

add_heading(doc, "2. Exclusions")
add_para(doc, "Structural framing, sheathing, and substrate preparation by others\nAny work not expressly included in the agreed scope and Schedule of Values above")

add_heading(doc, "3. Contract Amount and Payment")
add_para(doc, "Total Contract Amount: $91,650.00 (Ninety-One Thousand Six Hundred Fifty Dollars).")
add_para(doc, "Deposit: A deposit is required to reserve time, generally one (1) month prior to start. The percentage of deposit shall be agreed upon by both parties (typically 50/50 to cover travel, one-use tools and materials, and expenses associated with beginning the job). The deposit may be requested sooner depending on material chosen and lead times. Upon award of the bid and scheduling of the deposit, Subcontractor shall provide General Liability, Umbrella, and Workers’ Compensation certificates.")
add_para(doc, "Subcontractor shall submit progress invoices as agreed (e.g., monthly or by phase). Each invoice shall describe work completed and applicable percentage of contract value. General Contractor shall review, approve, and pay within thirty (30) days of receipt. A retention percentage may be agreed in writing; if not specified, ten percent (10%) may be withheld until final completion and punchlist sign-off.")
add_para(doc, "Invoices shall be submitted in writing (email acceptable) to: racinc19@gmail.com. Include: (1) invoice number and date, (2) billing period, (3) description of work completed, (4) percentage complete, (5) amount billed this period, (6) total billed to date.")

add_heading(doc, "4. Schedule")
add_para(doc, "Subcontractor shall commence work upon written Notice to Proceed from the General Contractor and upon confirmation of contract amount. Subcontractor shall maintain continuous progress as coordinated with the project schedule. Any delays caused by the Subcontractor may result in liquidated or actual damages at the General Contractor’s election.")

add_heading(doc, "5. Changes and Change Orders")
add_para(doc, "No extra work or changes to the Scope of Work shall be performed without prior written authorization from the General Contractor. Work performed without written approval will not be compensated. Subcontractor shall submit a written cost proposal for any change within five (5) business days of request.")

add_heading(doc, "6. Site Conditions and Responsibilities")
add_para(doc, "Subcontractor shall verify substrate and dimensions in the field. Discrepancies must be reported in writing immediately. Daily cleanup required. Subcontractor shall comply with all applicable OSHA regulations and maintain a safe worksite.")

add_heading(doc, "7. Insurance Requirements")
add_para(doc, "Prior to commencing work, Subcontractor shall obtain and maintain: Commercial General Liability: $1,000,000 per occurrence / $2,000,000 aggregate; Workers’ Compensation as required by Utah state law. RAC Inc. and ARJR Kanarraville Holdings, LLC shall be named as additional insureds on the CGL policy. Certificates of Insurance must be delivered to General Contractor before mobilization.")

add_heading(doc, "8. Indemnification")
add_para(doc, "Subcontractor shall indemnify, defend, and hold harmless RAC Inc., ARJR Kanarraville Holdings, LLC, and their officers, employees, and agents from and against any and all claims, damages, losses, and expenses, including attorney’s fees, arising out of or resulting from the Subcontractor’s performance under this Agreement, to the extent caused by the negligent acts or omissions of the Subcontractor or its employees.")

add_heading(doc, "9. Termination")
add_para(doc, "General Contractor may terminate for cause upon written notice if Subcontractor: (a) abandons the work; (b) fails to maintain required insurance; (c) performs defective work and fails to remedy within five (5) days of notice; or (d) becomes insolvent. General Contractor may terminate for convenience upon seven (7) days’ written notice, in which case Subcontractor shall be compensated for work satisfactorily completed to the date of termination.")

add_heading(doc, "10. Subcontractor Login Information (Project Master List)")
add_para(doc, "Project vendor and login access are tracked by Category, Vendor, and Pin per the project master list. General Contractor shall use this information only for project administration, document access, and communication and shall keep it confidential.")
doc.add_paragraph("Category: Rammed Earth")
doc.add_paragraph("Vendor: Alysa Oliver LLC (Allysa Oliver)")
doc.add_paragraph("Pin: 5550")
doc.add_paragraph()

doc.add_paragraph("Contract Date | _____________, ______")
doc.add_paragraph("Owner | ARJR Kanarraville Holdings, LLC")
doc.add_paragraph("General Contractor | RAC Inc.  |  Craig Engel  |  1267 N. 1390 W., St. George, UT 84770  |  435-229-7360")
doc.add_paragraph("Subcontractor | Allysa Oliver LLC  |  Allysa Oliver  |  hello@oliverfinishing.com")
doc.add_paragraph("Contract Amount | $91,650.00 (Ninety-One Thousand Six Hundred Fifty Dollars)  |  Estimate No. 60 (02/20/2026)")
doc.add_paragraph("Start Date | Upon NTP and deposit as agreed")
doc.add_paragraph()

doc.add_paragraph("GENERAL CONTRACTOR:")
doc.add_paragraph("RAC Inc.")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Craig Engel")
doc.add_paragraph("Title: General Contractor")
doc.add_paragraph("Date: _____________________________")

doc.add_paragraph()
doc.add_paragraph("SUBCONTRACTOR:")
doc.add_paragraph("Allysa Oliver LLC")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Allysa Oliver")
doc.add_paragraph("Title: _____________________________")
doc.add_paragraph("Date: _____________________________")

out = "Oliver_Finishing_Clayworks_Subcontract_Rodriguez_v2.docx"
doc.save(out)
print("Saved:", out)
