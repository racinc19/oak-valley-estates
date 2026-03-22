"""Generate CO-010 Addendum — ADU Shower Correction (body spray to non-body-spray)."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Header
p = doc.add_paragraph()
r = p.add_run('RAC INCORPORATED')
r.bold = True
r.font.size = Pt(32)

p = doc.add_paragraph('Craig Engel, General Contractor')
p.runs[0].font.size = Pt(19)

p = doc.add_paragraph('1267 N 1390 W, Saint George, UT 84770   |   (435) 229-7360   |   racinc19@gmail.com')
p.runs[0].font.size = Pt(17)

doc.add_paragraph()

# Title
p = doc.add_paragraph()
r = p.add_run('CHANGE ORDER CO-010 — ADDENDUM')
r.bold = True
r.font.size = Pt(36)
p = doc.add_paragraph()
r = p.add_run('ADU Shower Correction (Body Spray to Non–Body Spray)')
r.font.size = Pt(26)
p.paragraph_format.space_after = Pt(40)

# Info table
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
rows = table.rows
rows[0].cells[0].text = 'Project:'
rows[0].cells[1].text = 'Rodriguez Residence – Oak Valley Estates Lot 5'
rows[0].cells[2].text = 'Date:'
rows[0].cells[3].text = 'March 19, 2026'
rows[1].cells[0].text = 'Owner:'
rows[1].cells[1].text = 'Mr. & Mrs. Rodriguez'
rows[1].cells[2].text = 'CO Number:'
rows[1].cells[3].text = 'CO-010 (Addendum)'
rows[2].cells[0].text = 'Location:'
rows[2].cells[1].text = 'Liberty Lane, Kanarraville, UT 84742'
rows[2].cells[2].text = 'Credit:'
rows[2].cells[3].text = '($201.00)'
rows[3].cells[0].text = ''
rows[3].cells[1].text = ''
rows[3].cells[2].text = ''
rows[3].cells[3].text = ''

doc.add_paragraph()
doc.add_paragraph()

# Cost impact
p = doc.add_paragraph()
r = p.add_run('COST IMPACT: This addendum provides a $201.00 credit to the Owner. The ADU shower is changed from the body spray set to the non–body spray set. Net adjustment to CO-010: original $1,697.00 less $201.00 credit = $1,496.00.')
r.bold = True
r.font.size = Pt(14)
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(40)

# Scope
p = doc.add_paragraph()
r = p.add_run('SCOPE')
r.bold = True
r.font.size = Pt(24)
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph('ADU (Casita) shower fixture change:')
p.runs[0].font.size = Pt(12)
doc.add_paragraph()

# Detail table
detail = doc.add_table(rows=4, cols=4)
detail.style = 'Table Grid'
detail.rows[0].cells[0].text = 'Description'
detail.rows[0].cells[1].text = 'Model / SKU'
detail.rows[0].cells[2].text = 'Price'
detail.rows[0].cells[3].text = 'Notes'
detail.rows[1].cells[0].text = 'Originally specified (body spray set)'
detail.rows[1].cells[1].text = 'SGF10GJ-64-BL16'
detail.rows[1].cells[2].text = '$710.00'
detail.rows[1].cells[3].text = 'Rain head, secondary head, 6 body jets, handheld, 5-knob control'
detail.rows[2].cells[0].text = 'Revised to (non–body spray set)'
detail.rows[2].cells[1].text = 'SGF10GJ-03A-BL12'
detail.rows[2].cells[2].text = '$509.00'
detail.rows[2].cells[3].text = '12" triple handle 7-spray, rain head, handheld, anti-scald valve'
detail.rows[3].cells[0].text = 'Credit to Owner'
detail.rows[3].cells[1].text = ''
detail.rows[3].cells[2].text = '($201.00)'
detail.rows[3].cells[3].text = 'Difference credited'

doc.add_paragraph()
doc.add_paragraph()

# Authorization
p = doc.add_paragraph()
r = p.add_run('AUTHORIZATION')
r.bold = True
r.font.size = Pt(24)
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph('By signing below, Owner acknowledges the ADU shower change described above and accepts the $201.00 credit. This addendum amends Change Order CO-010. The revised net addition to the contract sum for CO-010 is $1,496.00.')
p.runs[0].font.size = Pt(12)
p.paragraph_format.space_after = Pt(30)

sig = doc.add_table(rows=3, cols=3)
sig.rows[0].cells[0].text = '___________________________'
sig.rows[0].cells[2].text = '___________________________'
sig.rows[1].cells[0].text = 'Craig Engel — RAC Incorporated'
sig.rows[1].cells[2].text = 'Owner — Rodriguez Residence'
sig.rows[2].cells[0].text = 'General Contractor'
sig.rows[2].cells[2].text = 'Date: ___________________'

# Save — same folder as CO-010 (script lives in Construction/Project/Change Order/CO-010/)
repo_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))))
out_path = os.path.join(repo_root, 'Construction', 'Change Order', 'CO-010', 'CO-010_Addendum_ADU_Shower_Correction.docx')
os.makedirs(os.path.dirname(out_path), exist_ok=True)
try:
    doc.save(out_path)
    print(f'Saved: {out_path}')
except PermissionError:
    alt = out_path.replace('.docx', '_new.docx')
    doc.save(alt)
    print(f'File may be open — saved to: {alt}')
