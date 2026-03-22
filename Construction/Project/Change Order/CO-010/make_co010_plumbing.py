"""Generate CO-010 Plumbing Fixture Changes change order."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('RAC INCORPORATED')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = None  # 1A2E4A - will use default

p = doc.add_paragraph('Craig Engel, General Contractor')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.runs[0].font.size = Pt(19)
p.runs[0].font.color.rgb = None

p = doc.add_paragraph('1267 N 1390 W, Saint George, UT 84770   |   (435) 229-7360   |   racinc19@gmail.com')
p.runs[0].font.size = Pt(17)
p.runs[0].font.color.rgb = None

doc.add_paragraph()

# Title
p = doc.add_paragraph()
r = p.add_run('CHANGE ORDER CO-010')
r.bold = True
r.font.size = Pt(36)
p = doc.add_paragraph()
r = p.add_run('Plumbing Fixture Changes')
r.font.size = Pt(26)
r.font.color.rgb = None
p.paragraph_format.space_after = Pt(40)

# Info table
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
rows = table.rows
rows[0].cells[0].text = 'Project:'
rows[0].cells[1].text = 'Rodriguez Residence – Oak Valley Estates Lot 5'
rows[0].cells[2].text = 'Date:'
rows[0].cells[3].text = 'March 19, 2026'
rows[1].cells[0].text = 'Owner:'
rows[1].cells[1].text = 'Mr. & Mrs. Rodriguez'
rows[1].cells[2].text = 'CO Number:'
rows[1].cells[3].text = 'CO-010'
rows[2].cells[0].text = 'Location:'
rows[2].cells[1].text = 'Liberty Lane, Kanarraville, UT 84742'
rows[2].cells[2].text = 'Contract Value:'
rows[2].cells[3].text = '$1,697.00'
rows[3].cells[0].text = ''
rows[3].cells[1].text = ''
rows[3].cells[2].text = ''
rows[3].cells[3].text = ''

doc.add_paragraph()
doc.add_paragraph()

# Cost summary box
p = doc.add_paragraph()
r = p.add_run('COST IMPACT: This change order adds $1,697.00 to the contract sum for plumbing fixture changes.')
r.bold = True
r.font.size = Pt(19)
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(40)

# Scope
p = doc.add_paragraph()
r = p.add_run('SCOPE SUMMARY')
r.bold = True
r.font.size = Pt(24)
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(20)

# Scope table
scope_table = doc.add_table(rows=2, cols=4)
scope_table.style = 'Table Grid'
scope_table.rows[0].cells[0].text = 'Item'
scope_table.rows[0].cells[1].text = 'Description'
scope_table.rows[0].cells[2].text = 'Cost'
scope_table.rows[0].cells[3].text = 'Total'
scope_table.rows[1].cells[0].text = 'Plumbing Fixture Changes'
scope_table.rows[1].cells[1].text = 'Owner-initiated changes to plumbing fixtures per revised selections'
scope_table.rows[1].cells[2].text = '$1,697.00'
scope_table.rows[1].cells[3].text = '$1,697.00'

doc.add_paragraph()
doc.add_paragraph()

# Cost breakdown
p = doc.add_paragraph()
r = p.add_run('COST BREAKDOWN')
r.bold = True
r.font.size = Pt(24)
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(20)

cost_table = doc.add_table(rows=3, cols=2)
cost_table.style = 'Table Grid'
cost_table.rows[0].cells[0].text = 'Item'
cost_table.rows[0].cells[1].text = 'Amount'
cost_table.rows[1].cells[0].text = 'Plumbing fixture changes'
cost_table.rows[1].cells[1].text = '$1,697.00'
cost_table.rows[2].cells[0].text = 'TOTAL CHANGE ORDER'
cost_table.rows[2].cells[1].text = '$1,697.00'

doc.add_paragraph()
doc.add_paragraph()

# Authorization
p = doc.add_paragraph()
r = p.add_run('AUTHORIZATION & OWNER APPROVAL')
r.bold = True
r.font.size = Pt(24)
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(20)

p = doc.add_paragraph('By signing below, Owner authorizes RAC Incorporated to proceed with the plumbing fixture changes described in Change Order CO-010. This change order adds $1,697.00 to the contract sum. This change order becomes part of the contract documents upon execution.')
p.runs[0].font.size = Pt(19)
p.paragraph_format.space_after = Pt(40)

# Signature table
sig_table = doc.add_table(rows=3, cols=3)
sig_table.rows[0].cells[0].text = '___________________________'
sig_table.rows[0].cells[2].text = '___________________________'
sig_table.rows[1].cells[0].text = 'Craig Engel — RAC Incorporated'
sig_table.rows[1].cells[2].text = 'Owner — Rodriguez Residence'
sig_table.rows[2].cells[0].text = 'General Contractor'
sig_table.rows[2].cells[2].text = 'Date: ___________________'

# Save
import os
out_path = r'c:\Users\racin\OneDrive\Desktop\Rodriquez\Construction\Change Order\CO-010\CO-010_Plumbing_Fixture_Changes.docx'
os.makedirs(os.path.dirname(out_path), exist_ok=True)
try:
    doc.save(out_path)
    print(f'Saved: {out_path}')
except PermissionError:
    alt = out_path.replace('.docx', '_updated.docx')
    doc.save(alt)
    print(f'File may be open - saved to: {alt}')
