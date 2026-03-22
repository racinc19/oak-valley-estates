# Create Word (.docx) for Blinds RFP — run from Bids\Blinds folder
from docx import Document
from docx.shared import Pt

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    return p

def main():
    import os
    folder = os.path.dirname(os.path.abspath(__file__))

    doc = Document()
    doc.add_paragraph('Request for Proposal', style='Title')
    doc.add_paragraph('Window Coverings — Blinds', style='Heading 1')
    doc.add_paragraph()
    add_para(doc, 'Project: Rodriguez Residence (Oak Valley Estates #5)')
    add_para(doc, 'Location: Oak Valley Estates Lot 5, 646 W Liberty Lane, Kanarraville, UT 84742')
    add_para(doc, 'Delivery: Lump sum or unit pricing per opening as specified below')
    add_para(doc, 'Reference: Attached A504 - Glass (Single Page). This is a bid — nothing is built yet; price from the drawing only. No field verification before bid.')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('Suggested scope (per attached A504 - Glass)', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('Motorized blinds ').bold = True
    p.add_run('(supply motors and blinds) for all glazed openings indicated on the drawing:')
    for item in ['2-Main Glass', '3-Upper Glass', 'Glass East', 'Glass South', 'Glass West', 'Glass West 2', 'Glass Entry', 'Glass Entry 2']:
        doc.add_paragraph(item, style='List Bullet')
    add_para(doc, 'Review A504 - Glass and base your bid on the openings shown. State any exclusions (e.g. fixed glass, doors) and pricing impact.')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('1. Scope of Work', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('Supply and install ').bold = False
    p.add_run('motorized blinds').bold = True
    p.add_run(' (motors and blinds) at all openings per suggested scope above. Scope includes:')
    doc.add_paragraph('Furnish and install motorized blinds at all windows/glass indicated on A504 - Glass. Dimensions and counts from the drawing; no field measurement before bid.', style='List Bullet')
    doc.add_paragraph('Installation: appropriate mounting per opening; all hardware, brackets, headrails, and motors included.', style='List Bullet')
    doc.add_paragraph('Removal of packaging and disposal; leave openings clean and operational.', style='List Bullet')
    doc.add_paragraph('Installation to be scheduled after interior trim and paint are complete.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('2. What to Include in Your Bid (material, blackout, etc.)', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('We need ').bold = False
    p.add_run('one clear price').bold = True
    p.add_run(' from the drawing. To avoid back-and-forth:')
    doc.add_paragraph("State exactly what you're bidding: One material type (e.g. cellular, roller, faux-wood — your standard). One opacity (e.g. light filtering; or state \"light filtering standard, blackout available as alternate\").", style='List Bullet')
    doc.add_paragraph('One base lump sum or unit price for that scope. If blackout (or different material) changes the price, give it as a separate line item or alternate (e.g. "Blackout add $___ per opening" or "Alternate: blackout throughout $___").', style='List Bullet')
    doc.add_paragraph('Color/finish: selected by Owner after award; state options in the proposal.', style='List Bullet')
    add_para(doc, "We'll decide material/blackout once we have your number. Your bid shall clearly state: product line, what's included in the base price (material + opacity), and any alternates.")
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('3. Schedule', style='Heading 2')
    doc.add_paragraph('Lead time from order to delivery and installation (measurement will be scheduled once trim is in place).', style='List Bullet')
    doc.add_paragraph('Installation to be coordinated with GC after interior trim and paint; state estimated install duration.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('4. Pricing', style='Heading 2')
    doc.add_paragraph('Lump sum for full scope (all openings per A504 - Glass), or unit price per opening with estimated total. Base your count on the drawing.', style='List Bullet')
    doc.add_paragraph('Price includes: product (blinds + motors), hardware, installation, removal of debris.', style='List Bullet')
    doc.add_paragraph('Alternates as separate line items if you\'re quoting more than one option (e.g. blackout, different material).', style='List Bullet')
    doc.add_paragraph('Payment terms (e.g. deposit on order, balance on completion; or per GC draw schedule).', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('5. Submittals', style='Heading 2')
    doc.add_paragraph('Product data sheets / cut sheets for proposed blind type(s).', style='List Bullet')
    doc.add_paragraph('Sample or sample availability for color/finish selection (post-award if needed).', style='List Bullet')
    doc.add_paragraph('Warranty (manufacturer and installer).', style='List Bullet')
    doc.add_paragraph('Exclusions or assumptions (e.g. specialty shapes, skylights, doors—included or excluded).', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('6. Proposal Format and Due Date', style='Heading 2')
    doc.add_paragraph('Cover letter with company name, contact person, phone, and email.', style='List Number')
    doc.add_paragraph('Scope summary: product type(s), what\'s in your base price (material + opacity), number of openings per drawing, alternates if any.', style='List Number')
    doc.add_paragraph('Pricing: lump sum and/or unit breakdown; payment terms.', style='List Number')
    doc.add_paragraph('Lead time and availability for order and installation.', style='List Number')
    doc.add_paragraph('Exceptions or assumptions.', style='List Number')
    add_para(doc, 'Submit proposals to: [Insert contact name/email]')
    add_para(doc, 'Due date: [Insert date]')
    doc.add_paragraph()
    add_para(doc, 'This RFP is a bid from the drawing only (no field verification before bid). Scope: supply and install motorized blinds (motors + blinds) per A504 - Glass. Contractor states in the proposal one base price and what it includes (material, opacity); alternates (e.g. blackout) as separate line items. Color/finish selected by Owner after award.', bold=False)
    out = os.path.join(folder, 'RFP - Blinds.docx')
    doc.save(out)
    print('Created:', out)

if __name__ == '__main__':
    main()
