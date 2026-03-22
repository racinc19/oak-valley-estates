"""Generate Rodriguez_Project_Supervision_Knudsen.docx from contract content."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Run: pip install python-docx")

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_para(doc, text, bullet=False):
    p = doc.add_paragraph(text, style="List Bullet" if bullet else "Normal")
    p.paragraph_format.space_after = Pt(4)

doc = Document()
doc.add_paragraph("CONTRACT — PROJECT SUPERVISION").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence | Oak Valley Estates").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "PARTIES")
doc.add_paragraph("Company:    RAC Development")
doc.add_paragraph("            [Address]")
doc.add_paragraph("            Craig Engel | (435) 229-7360 | racinc19@gmail.com")
doc.add_paragraph()
doc.add_paragraph("Contractor: Daniel Knudsen")
doc.add_paragraph("            [Address]")
doc.add_paragraph("            [Phone]")
doc.add_paragraph("            [Email]")

add_heading(doc, "PROJECT & ROLE")
doc.add_paragraph("Project:    Rodriguez Residence job site (Oak Valley Estates, Kanarraville, UT).")
doc.add_paragraph("Scope:      On-site project supervision per daily direction from RAC. RAC retains project management; Contractor provides field supervision and presence as required.")

add_heading(doc, "SCHEDULE & HOURS")
add_para(doc, "Report time: 8:00 AM at the Rodriguez job site on each day supervision is required.", bullet=True)
add_para(doc, "Concrete days: Report time may be earlier (approximately 2–3 hours earlier when concrete is pouring). RAC will notify Contractor the day before when early arrival is required.", bullet=True)
add_para(doc, "Daily duration: Approximately two (2) to three (3) hours per day, as needed.", bullet=True)
add_para(doc, "Days required: As directed by RAC. RAC will provide a daily list (or schedule) for each day that supervision is required.", bullet=True)

add_heading(doc, "COMPENSATION")
add_para(doc, "Rate: $30.00 per hour for all supervised hours on site.", bullet=True)
add_para(doc, "Travel: RAC will pay travel one (1) way per day when Contractor is required on site, at the same rate: $30.00 per hour for travel time (not mileage).", bullet=True)
add_para(doc, "Invoicing: Contractor shall submit invoices every Friday for hours worked and travel. Payment will be made by the following Friday.", bullet=True)

add_heading(doc, "RAC RESPONSIBILITIES")
add_para(doc, "RAC handles project management, scheduling, and trade coordination.", bullet=True)
add_para(doc, "RAC will provide Contractor with a daily list (scope/tasks) for each day supervision is required.", bullet=True)
add_para(doc, "RAC will notify Contractor of any early start (e.g., concrete) no later than the prior day.", bullet=True)

add_heading(doc, "CONTRACTOR RESPONSIBILITIES")
add_para(doc, "Arrive on site at the designated time (8:00 AM or earlier when notified for concrete).", bullet=True)
add_para(doc, "Perform field supervision per the daily list provided by RAC.", bullet=True)
add_para(doc, "Communicate promptly with RAC regarding site conditions, issues, or delays.", bullet=True)
add_para(doc, "Maintain professional conduct and safe practices on site.", bullet=True)

add_heading(doc, "TERM & TERMINATION")
add_para(doc, "Term: This agreement is effective as of [Effective Date]. The job duration is approximately one (1) year, or until the Rodriguez project no longer requires supervision, whichever is earlier.", bullet=True)
add_para(doc, "Termination: Either party may be excused from this contract with one (1) week's (7 days) written notice. RAC will pay for all hours worked and approved travel through the termination date.", bullet=True)

add_heading(doc, "STATUS & TAXES")
add_para(doc, "Contractor is engaged as an independent contractor, not an employee. Contractor is responsible for his own taxes, insurance, and benefits. RAC will not withhold taxes or provide benefits.", bullet=True)

add_heading(doc, "LIABILITY & INSURANCE")
add_para(doc, "Contractor shall maintain general liability insurance appropriate for on-site supervision and shall provide proof of insurance upon request.", bullet=True)
add_para(doc, "Contractor shall not hold RAC liable for injury or loss arising from Contractor's own negligence. RAC shall not be liable for Contractor's tools, vehicle, or personal property.", bullet=True)

add_heading(doc, "OTHER")
add_para(doc, "Confidentiality: Contractor shall not disclose project details, pricing, or proprietary information to third parties without RAC's consent.", bullet=True)
add_para(doc, "Assignment: Contractor may not assign this agreement without RAC's written consent.", bullet=True)
add_para(doc, "Entire agreement: This document constitutes the full agreement. Changes must be in writing and signed by both parties.", bullet=True)
add_para(doc, "Governing law: Laws of the State of Utah.", bullet=True)

add_heading(doc, "SIGNATURES")
doc.add_paragraph("RAC Development")
doc.add_paragraph()
doc.add_paragraph("By: _______________________________   Date: _______________")
doc.add_paragraph("    (Authorized signatory)")
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph("Daniel Knudsen")
doc.add_paragraph()
doc.add_paragraph("By: _______________________________   Date: _______________")

out = Path(__file__).parent / "Rodriguez_Project_Supervision_Knudsen_v2.docx"
doc.save(out)
print("Saved:", out)
