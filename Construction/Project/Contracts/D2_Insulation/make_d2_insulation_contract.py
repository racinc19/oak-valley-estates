# Creates D2 Building Solutions Insulation Subcontract (Rodriguez Residence)
# Run: python make_d2_insulation_contract.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()
doc.add_paragraph("SUBCONTRACT AGREEMENT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("INSULATION").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence  |  Oak Valley Estates Lot 5  |  Kanarraville, UT 84742")
doc.add_paragraph()

add_heading(doc, "1. Scope of Work")
add_para(doc, "Subcontractor shall furnish all labor and materials necessary to perform insulation work for the Rodriguez Residence (Oak Valley Estates Lot 5, Liberty Lane, Kanarraville) per Estimate No. 802230 and the Contract Documents. Base scope includes:")
add_para(doc, "• Ceiling area: Open cell foam R-49\n• Rim joist: Open cell foam R-30\n• Exterior walls: R-23 mesh & blow fiberglass – high performance insulation system\n• Foam – windows, doors, penetrations and cracks\n• Basement furred walls: R-15 mesh & blow fiberglass – high performance insulation system\n• Fire blocking: R-15 Rockwool batt\n• Garage ceiling area: Open cell foam R-30\n• Garage exterior walls: R-23 mesh & blow fiberglass – high performance insulation system")
add_para(doc, "Base Estimate Total: $36,778.00 (Thirty-Six Thousand Seven Hundred Seventy-Eight Dollars). Additional options (e.g., between-floors R-38 net & blow cellulose, sound walls) may be added by written change order at quoted prices.")
add_para(doc, "Spray foam is applied based on nominal thickness; open cell acceptable variance is ±1/2 inch per industry standards. Final performance is based on average application thickness; localized deviations within the accepted range do not constitute a defect.")

add_heading(doc, "2. Exclusions")
add_para(doc, "Any work not specifically described in this agreement or in an approved change order shall be considered extra and requires written authorization from the General Contractor prior to commencement.")

add_heading(doc, "3. Contract Amount and Payment")
add_para(doc, "Base Contract Amount: $36,778.00 (Thirty-Six Thousand Seven Hundred Seventy-Eight Dollars).")
add_para(doc, "Subcontractor shall submit invoices as agreed. General Contractor shall pay within thirty (30) days of receipt. Past-due amounts may accrue interest at 1.5% per month until paid in full per Subcontractor’s policy. A retention percentage may be agreed in writing; if not specified, ten percent (10%) may be withheld until final completion and punchlist sign-off.")
add_para(doc, "Invoices shall be submitted in writing (email acceptable) to: racinc19@gmail.com. Include: (1) invoice number and date, (2) billing period, (3) description of work completed, (4) percentage complete, (5) amount billed this period, (6) total billed to date.")

add_heading(doc, "4. Schedule")
add_para(doc, "Subcontractor shall commence work upon written Notice to Proceed from the General Contractor. Subcontractor shall maintain continuous progress as coordinated with the project schedule. Any delays caused by the Subcontractor may result in liquidated or actual damages at the General Contractor’s election. Quote valid for 30 days from estimate date; contract execution confirms pricing.")

add_heading(doc, "5. Changes and Change Orders")
add_para(doc, "No extra work or changes to the Scope of Work shall be performed without prior written authorization from the General Contractor. Alterations or deviations involving extra costs shall be executed upon written (or verbal, per Subcontractor’s practices) change order and shall be an extra charge. Subcontractor may adjust quoted prices in the event of shortages, environmental impacts, freight surcharges, manufacturer price increases, or regulations; such adjustments require written notice.")

add_heading(doc, "6. Site Conditions and Responsibilities")
add_para(doc, "Subcontractor shall verify substrate and conditions in the field. Discrepancies must be reported in writing immediately. Daily cleanup required. Subcontractor shall comply with all applicable OSHA regulations and maintain a safe worksite.")

add_heading(doc, "7. Insurance Requirements")
add_para(doc, "Prior to commencing work, Subcontractor shall obtain and maintain: Commercial General Liability: $1,000,000 per occurrence / $2,000,000 aggregate; Workers’ Compensation as required by Utah state law. RAC Inc. and ARJR Kanarraville Holdings, LLC shall be named as additional insureds on the CGL policy. Certificates of Insurance must be delivered to General Contractor before mobilization.")

add_heading(doc, "8. Indemnification")
add_para(doc, "Subcontractor shall indemnify, defend, and hold harmless RAC Inc., ARJR Kanarraville Holdings, LLC, and their officers, employees, and agents from and against any and all claims, damages, losses, and expenses, including attorney’s fees, arising out of or resulting from the Subcontractor’s performance under this Agreement, to the extent caused by the negligent acts or omissions of the Subcontractor or its employees.")

add_heading(doc, "9. Termination")
add_para(doc, "General Contractor may terminate for cause upon written notice if Subcontractor: (a) abandons the work; (b) fails to maintain required insurance; (c) performs defective work and fails to remedy within five (5) days of notice; or (d) becomes insolvent. General Contractor may terminate for convenience upon seven (7) days’ written notice, in which case Subcontractor shall be compensated for work satisfactorily completed to the date of termination.")

add_heading(doc, "10. Subcontractor Login Information (Project Master List)")
add_para(doc, "Project vendor and login access are tracked by Category, Vendor, and Pin per the project master list. General Contractor shall use this information only for project administration, document access, and communication and shall keep it confidential.")
doc.add_paragraph("Category: Insulation")
doc.add_paragraph("Vendor: D2 Building Solutions")
doc.add_paragraph("Pin: _________________________")
doc.add_paragraph()

doc.add_paragraph("Contract Date | _____________, ______")
doc.add_paragraph("Owner | ARJR Kanarraville Holdings, LLC")
doc.add_paragraph("General Contractor | RAC Inc.  |  Craig Engel  |  1267 N. 1390 W., St. George, UT 84770  |  435-229-7360")
doc.add_paragraph("Subcontractor | D2 Building Solutions  |  5099 Wheeler Way, Hurricane, UT 84737  |  (435) 688-1337  |  nathan@d2bldg.com")
doc.add_paragraph("Contract Amount | $36,778.00 (Thirty-Six Thousand Seven Hundred Seventy-Eight Dollars)  |  Estimate No. 802230")
doc.add_paragraph("Start Date | Upon NTP")
doc.add_paragraph()

doc.add_paragraph("GENERAL CONTRACTOR:")
doc.add_paragraph("RAC Inc.")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Craig Engel")
doc.add_paragraph("Title: General Contractor")
doc.add_paragraph("Date: _____________________________")

doc.add_paragraph()
doc.add_paragraph("SUBCONTRACTOR:")
doc.add_paragraph("D2 Building Solutions")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: _____________________________")
doc.add_paragraph("Title: _____________________________")
doc.add_paragraph("Date: _____________________________")

out = "D2_Insulation_Subcontract_Rodriguez.docx"
doc.save(out)
print("Saved:", out)
