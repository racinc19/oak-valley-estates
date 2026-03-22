# Creates H & E Drywall Subcontract (Rodriguez Residence)
# Run: python make_he_drywall_contract.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()
doc.add_paragraph("SUBCONTRACT AGREEMENT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("DRYWALL / SHEETROCK").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence  |  Oak Valley Estates Lot 5  |  Liberty Lane, Kanarraville, UT 84742")
doc.add_paragraph()

add_heading(doc, "1. Scope of Work")
add_para(doc, "Subcontractor shall furnish all labor, materials, and supplies necessary to perform drywall work for the Rodriguez Residence (Oak Valley Estates Lot 5, Liberty Lane, Kanarraville) per the proposal dated 09/02/2025 and the Contract Documents. Scope includes:")
add_para(doc, "Drywall: 1/2\" drywall installation in garage ceilings and walls; 1/2\" drywall in main floor ceilings and walls; 1/2\" drywall in second floor ceilings and walls; 1/2\" drywall in basement ceilings and walls. 1/2\" Densshield tile backer installation in showers. 5/8\" exterior board in porch ceilings.")
add_para(doc, "Corners: Square or round per preference; no kerf doors; 4-way window wrap.")
add_para(doc, "Taping: Three-coat system to prepare drywall for texture.")
add_para(doc, "Texture: Old World texture on walls (holy smooth) and Old World texture on ceilings (holy smooth). Texture is varied and subjective; if approval by others is required, Subcontractor must be notified before texturing begins so it can be approved with them. Subcontractor is not responsible for retexturing or paint changes if someone is dissatisfied after texturing is completed.")
add_para(doc, "Clean-up: Garage floor masked off to protect from drywall mud; complete clean-up of drywall items; floors scraped and swept.")
add_para(doc, "Repairs: One (1) final trip to fix nicks and dings; no more than three (3) small patches included. Additional repairs, repairs resulting from poor framing/bowed walls/uneven trusses, or work for windows/pocket doors/framing added after hanging has begun will be billed as separate or extra expense.")
add_para(doc, "Total Contract Amount: $41,400.00 (Forty-One Thousand Four Hundred Dollars). Estimates include all labor, materials, and supplies.")

add_heading(doc, "2. Exclusions")
add_para(doc, "Repairs beyond one final trip and three small patches (billed separately). Retexturing or paint changes after completion. Damages to doors or tubs not covered/protected by others. Materials for other trades should not be delivered while Subcontractor is working; Subcontractor is not responsible for damages. Any work not specifically described in this agreement requires written authorization from the General Contractor prior to commencement.")

add_heading(doc, "3. Contract Amount and Payment")
add_para(doc, "Total Contract Amount: $41,400.00 (Forty-One Thousand Four Hundred Dollars).")
add_para(doc, "Invoices shall be paid in full within thirty (30) days after texture is complete. Late payments may incur a charge of 2% per month. A retention percentage may be agreed in writing; if not specified, ten percent (10%) may be withheld until final completion and punchlist sign-off.")
add_para(doc, "Invoices shall be submitted in writing (email acceptable) to: racinc19@gmail.com. Include: (1) invoice number and date, (2) billing period, (3) description of work completed, (4) percentage complete, (5) amount billed this period, (6) total billed to date.")

add_heading(doc, "4. Schedule")
add_para(doc, "Subcontractor shall commence work upon written Notice to Proceed from the General Contractor. Subcontractor shall maintain continuous progress as coordinated with the project schedule. Any delays caused by the Subcontractor may result in liquidated or actual damages at the General Contractor's election. This proposal is void if not accepted by signature within 30 days of the proposal date.")

add_heading(doc, "5. Changes and Change Orders")
add_para(doc, "No extra work or changes to the Scope of Work shall be performed without prior written authorization from the General Contractor. Work performed without written approval will not be compensated. Subcontractor shall submit a written cost proposal for any change within five (5) business days of request.")

add_heading(doc, "6. Site Conditions and Responsibilities")
add_para(doc, "Doors shall be covered and protected by others; tubs shall be covered with a hard walkable surface. Subcontractor shall verify substrate and dimensions in the field. Discrepancies must be reported in writing immediately. Subcontractor shall comply with all applicable OSHA regulations and maintain a safe worksite.")

add_heading(doc, "7. Insurance Requirements")
add_para(doc, "Prior to commencing work, Subcontractor shall obtain and maintain: Commercial General Liability: $1,000,000 per occurrence / $2,000,000 aggregate; Workers' Compensation as required by Utah state law (Subcontractor represents workers are fully covered). RAC Inc. and ARJR Kanarraville Holdings, LLC shall be named as additional insureds on the CGL policy. Certificates of Insurance must be delivered to General Contractor before mobilization.")

add_heading(doc, "8. Indemnification")
add_para(doc, "Subcontractor shall indemnify, defend, and hold harmless RAC Inc., ARJR Kanarraville Holdings, LLC, and their officers, employees, and agents from and against any and all claims, damages, losses, and expenses, including attorney's fees, arising out of or resulting from the Subcontractor's performance under this Agreement, to the extent caused by the negligent acts or omissions of the Subcontractor or its employees.")

add_heading(doc, "9. Termination")
add_para(doc, "General Contractor may terminate for cause upon written notice if Subcontractor: (a) abandons the work; (b) fails to maintain required insurance; (c) performs defective work and fails to remedy within five (5) days of notice; or (d) becomes insolvent. General Contractor may terminate for convenience upon seven (7) days' written notice, in which case Subcontractor shall be compensated for work satisfactorily completed to the date of termination.")

add_heading(doc, "10. Subcontractor Login Information (Project Master List)")
add_para(doc, "Project vendor and login access are tracked by Category, Vendor, and Pin per the project master list. General Contractor shall use this information only for project administration, document access, and communication and shall keep it confidential.")
doc.add_paragraph("Category: Sheetrock")
doc.add_paragraph("Vendor: H & E Drywall")
doc.add_paragraph("Pin: _________________________")
doc.add_paragraph()

doc.add_paragraph("Contract Date | _____________, ______")
doc.add_paragraph("Owner | ARJR Kanarraville Holdings, LLC")
doc.add_paragraph("General Contractor | RAC Inc.  |  Craig Engel  |  1267 N. 1390 W., St. George, UT 84770  |  435-229-7360")
doc.add_paragraph("Subcontractor | H & E Drywall  |  1018 S 1740 E, St. George, UT 84790  |  (435) 236-5652  |  hedrywall1022@gmail.com  |  UT Lic. #12226934-5501")
doc.add_paragraph("Contract Amount | $41,400.00 (Forty-One Thousand Four Hundred Dollars)  |  Proposal 09/02/2025")
doc.add_paragraph("Start Date | Upon NTP")
doc.add_paragraph()

doc.add_paragraph("GENERAL CONTRACTOR:")
doc.add_paragraph("RAC Inc.")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Craig Engel")
doc.add_paragraph("Title: General Contractor")
doc.add_paragraph("Date: _____________________________")

doc.add_paragraph()
doc.add_paragraph("SUBCONTRACTOR:")
doc.add_paragraph("H & E Drywall")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: _____________________________")
doc.add_paragraph("Title: _____________________________")
doc.add_paragraph("Date: _____________________________")

out = "H_E_Drywall_Subcontract_Rodriguez.docx"
doc.save(out)
print("Saved:", out)
