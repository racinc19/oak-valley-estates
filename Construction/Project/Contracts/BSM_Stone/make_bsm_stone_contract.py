# Creates BSM Residential Stone Veneer Subcontract (Rodriguez Residence)
# Run: python make_bsm_stone_contract.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

doc = Document()
doc.add_paragraph("SUBCONTRACT AGREEMENT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("STONE THIN VENEER").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Rodriguez Residence  |  Oak Valley Estates Lot 5  |  Kanarraville, UT 84742")
doc.add_paragraph()

add_heading(doc, "1. Scope of Work")
add_para(doc, "Subcontractor shall furnish all labor and materials necessary to perform exterior stone thin veneer installation for the Rodriguez Residence (Oak Valley Estates Lot 5) per Bid No. STG25-319 and the Contract Documents. Scope includes: Koni Fast Set Calabria, drystack application; mortar color poly modified; exterior stone thin veneer installation as detailed in the bid (quantities and rates per bid). Price is subject to change based on field measure and final stone selection; any such changes shall be documented by written change order.")
add_para(doc, "Schedule of Values (Bid STG25-319):")
add_para(doc, "• Exterior Calabria Fast_Set, mortar color poly modified, drystack: 160 sf @ $8.00 = $1,488.00\n• Exterior Calabria Fast_Set, mortar color poly modified, drystack: 1,700 sf @ $15.00 = $27,435.00\n• Project Total (Labor & Materials): $28,923.00")
add_para(doc, "Total Contract Amount: $28,923.00 (Twenty-Eight Thousand Nine Hundred Twenty-Three Dollars). Progressive billing may apply as agreed by the parties.")
add_para(doc, "Subcontractor’s masons are fully insured. Subcontractor provides a 1-year limited warranty on stone installation. Varying grout sizes, cut pieces, distribution of color tones, and seeping mortar (dry stack) are considered character of the stone and not installation faults. Warranty does not cover circumstances uncontrollable by the installer (e.g., heaving concrete, shifting/twisting of non-masonry columns, acts of God, excessive settlement, cracked corners on columns due to construction). BSM is not responsible for waterproofing; customer is recommended to caulk areas of concern (water edge/siding, stone meets windows, stone meets non-masonry material) upon completion.")

add_heading(doc, "2. Exclusions")
add_para(doc, "Tyvek; lintels; rebar; backer rod and caulk; metal flashing; dumpster fee; water; electricity; winterization. Lath and scratch/brown for stone prep excluded (by others unless otherwise agreed in writing). Any work not specifically described in this agreement requires written authorization from the General Contractor prior to commencement.")

add_heading(doc, "3. Contract Amount and Payment")
add_para(doc, "Total Contract Amount: $28,923.00 (Twenty-Eight Thousand Nine Hundred Twenty-Three Dollars).")
add_para(doc, "Subcontractor shall submit progress invoices as agreed (e.g., progressive billing by phase). Each invoice shall describe work completed and applicable percentage of contract value. General Contractor shall review, approve, and pay within thirty (30) days of receipt. A retention percentage may be agreed in writing; if not specified, ten percent (10%) may be withheld until final completion and punchlist sign-off. Past-due balances may accrue interest at 1.5% per month plus reasonable collection/attorney fees as stated in Subcontractor’s terms.")
add_para(doc, "Invoices shall be submitted in writing (email acceptable) to: racinc19@gmail.com. Include: (1) invoice number and date, (2) billing period, (3) description of work completed, (4) percentage complete, (5) amount billed this period, (6) total billed to date.")

add_heading(doc, "4. Schedule")
add_para(doc, "Subcontractor shall commence work upon written Notice to Proceed from the General Contractor. Subcontractor shall maintain continuous progress as coordinated with the project schedule. Any delays caused by the Subcontractor may result in liquidated or actual damages at the General Contractor’s election. Prices are void if contract is not signed and returned within 14 days of bid date.")

add_heading(doc, "5. Changes and Change Orders")
add_para(doc, "No extra work or changes to the Scope of Work shall be performed without prior written authorization from the General Contractor. Work performed without written approval will not be compensated. Subcontractor shall submit a written cost proposal for any change within five (5) business days of request. Field measure and final stone selection may result in a change order to the contract amount.")

add_heading(doc, "6. Site Conditions and Responsibilities")
add_para(doc, "Subcontractor shall verify substrate and dimensions in the field. Discrepancies must be reported in writing immediately. Daily cleanup required. Subcontractor shall comply with all applicable OSHA regulations and maintain a safe worksite.")

add_heading(doc, "7. Insurance Requirements")
add_para(doc, "Prior to commencing work, Subcontractor shall obtain and maintain: Commercial General Liability: $1,000,000 per occurrence / $2,000,000 aggregate; Workers’ Compensation as required by Utah state law. RAC Inc. and ARJR Kanarraville Holdings, LLC shall be named as additional insureds on the CGL policy. Certificates of Insurance must be delivered to General Contractor before mobilization.")

add_heading(doc, "8. Indemnification")
add_para(doc, "Subcontractor shall indemnify, defend, and hold harmless RAC Inc., ARJR Kanarraville Holdings, LLC, and their officers, employees, and agents from and against any and all claims, damages, losses, and expenses, including attorney’s fees, arising out of or resulting from the Subcontractor’s performance under this Agreement, to the extent caused by the negligent acts or omissions of the Subcontractor or its employees.")

add_heading(doc, "9. Termination")
add_para(doc, "General Contractor may terminate for cause upon written notice if Subcontractor: (a) abandons the work; (b) fails to maintain required insurance; (c) performs defective work and fails to remedy within five (5) days of notice; or (d) becomes insolvent. General Contractor may terminate for convenience upon seven (7) days’ written notice, in which case Subcontractor shall be compensated for work satisfactorily completed to the date of termination.")

add_heading(doc, "10. Subcontractor Login Information (Project Master List)")
add_para(doc, "Project vendor and login access are tracked by Category, Vendor, and Pin per the project master list. General Contractor shall use this information only for project administration, document access, and communication and shall keep it confidential.")
doc.add_paragraph("Category: Stone Veneer")
doc.add_paragraph("Vendor: BSM Residential (Builder's Stone & Masonry)")
doc.add_paragraph("Pin: _________________________")
doc.add_paragraph()

doc.add_paragraph("Contract Date | _____________, ______")
doc.add_paragraph("Owner | ARJR Kanarraville Holdings, LLC")
doc.add_paragraph("General Contractor | RAC Inc.  |  Craig Engel  |  1267 N. 1390 W., St. George, UT 84770  |  435-229-7360")
doc.add_paragraph("Subcontractor | BSM Residential  |  3376 E 380 St N, St. George, UT 84790  |  (435) 466-1100  |  www.bsmresidential.com  |  josette@bsmmasonry.com")
doc.add_paragraph("Contract Amount | $28,923.00 (Twenty-Eight Thousand Nine Hundred Twenty-Three Dollars)  |  Bid No. STG25-319")
doc.add_paragraph("Start Date | Upon NTP")
doc.add_paragraph()

doc.add_paragraph("GENERAL CONTRACTOR:")
doc.add_paragraph("RAC Inc.")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: Craig Engel")
doc.add_paragraph("Title: General Contractor")
doc.add_paragraph("Date: _____________________________")

doc.add_paragraph()
doc.add_paragraph("SUBCONTRACTOR:")
doc.add_paragraph("BSM Residential")
doc.add_paragraph("Signature: _______________________________")
doc.add_paragraph("Name: _____________________________")
doc.add_paragraph("Title: _____________________________")
doc.add_paragraph("Date: _____________________________")

out = "BSM_Stone_Subcontract_Rodriguez.docx"
doc.save(out)
print("Saved:", out)
