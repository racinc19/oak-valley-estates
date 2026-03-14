# Create Word (.docx) for Site Camera RFP — INSTALL ONLY — run from Site Camera folder
from docx import Document
from docx.shared import Pt

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    return p

def main():
    import os
    folder = os.path.dirname(os.path.abspath(__file__))

    doc = Document()
    doc.add_paragraph('Request for Proposal', style='Title')
    doc.add_paragraph('Site Camera System — Installation Only', style='Heading 1')
    doc.add_paragraph()
    add_para(doc, 'Project: Rodriguez Residence')
    add_para(doc, 'Location: Oak Valley Estates Lot 5, 646 W Liberty Lane, Kanarraville, UT 84742')
    add_para(doc, 'Delivery: Lump sum for installation labor per attached scope')
    add_para(doc, 'Reference: All equipment and design are Owner-furnished. See installer reference documents in this folder (e.g. 3_Starlink_Security_Visual.html) for full scope and sequence.')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('1. Scope of Work — Installation Only', style='Heading 2')
    add_para(doc, 'Proposals are for INSTALLATION ONLY. All equipment (cameras, poles, Starlink, solar, battery, brackets, concrete) is Owner-furnished. Contractor performs labor per the attached installer reference documents.', bold=False)
    doc.add_paragraph('Five (5) poles: North (Starlink + solar + camera), South/East/West (camera only). Layout per site diagram.', style='List Bullet')
    doc.add_paragraph('Concrete: 18" dia hole, 36" deep per pole; 2 bags 80 lb concrete; plumb; 24 hrs cure before mounting.', style='List Bullet')
    doc.add_paragraph('North pole: Starlink dish (top), Renogy 100W on LINOVISION bracket (30–35° south), aosu camera (MorningRo clamp, aim south), Starlink battery (5–6 ft). Wire battery before solar input. MC4 panel to battery; Starlink cable to battery.', style='List Bullet')
    doc.add_paragraph('South/East/West poles: aosu camera at 10–11 ft on MorningRo clamp. South→NORTH, East→WEST, West→EAST. Built-in solar cameras — no pole wiring.', style='List Bullet')
    doc.add_paragraph('Camera app setup: Power all 4 cameras; aosu app; scan QR; connect to Starlink WiFi (2.4 GHz); confirm live feed. Starlink must be live first.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('2. Owner-Furnished', style='Heading 2')
    doc.add_paragraph('All poles, hardware, Starlink (dish, battery), Renogy 100W, LINOVISION bracket, aosu cameras (4), MorningRo clamps, concrete, and materials per installer reference docs.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('3. Contractor Provides', style='Heading 2')
    doc.add_paragraph('Labor: layout, excavation, concrete, pole setting, mounting per reference, north-pole wiring, camera app setup and verification. Tools/equipment; state any beyond basic. Confirm Starlink and all 4 cameras live before demobilization.', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('4. Pricing', style='Heading 2')
    doc.add_paragraph('Lump sum for installation only. Payment terms. Exclusions or assumptions (access, staging, Owner presence for app).', style='List Bullet')
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('5. Proposal Format and Due Date', style='Heading 2')
    doc.add_paragraph('(1) Cover letter — company, contact. (2) Confirmation scope is install only per reference docs. (3) Lump sum, payment terms, exclusions. (4) Availability/lead time. (5) Questions.', style='List Number')
    add_para(doc, 'Submit to: [Insert contact name/email]')
    add_para(doc, 'Due date: [Insert date]')
    doc.add_paragraph()
    add_para(doc, 'This RFP is for installation labor only. Equipment and design are in the Site Camera folder reference documents (e.g. 3_Starlink_Security_Visual.html).', bold=False)
    out = os.path.join(folder, 'RFP - Site Camera Install.docx')
    doc.save(out)
    print('Created:', out)

if __name__ == '__main__':
    main()
